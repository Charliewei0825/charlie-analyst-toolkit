#!/usr/bin/env python3
"""Extract paragraphs and tables from DOCX files."""

from __future__ import annotations

import argparse
from collections import Counter
import json
from pathlib import Path
from typing import Any


def _clean(text: str | None) -> str:
    if text is None:
        return ""
    return text.replace("\r\n", "\n").strip()


def extract_docx_document(docx_path: Path) -> dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:
        raise SystemExit("Missing dependency `python-docx`. Install with: pip install python-docx") from exc

    doc = Document(str(docx_path))

    paragraphs: list[dict[str, Any]] = []
    for index, para in enumerate(doc.paragraphs, start=1):
        text = _clean(para.text)
        if not text:
            continue
        style_name = para.style.name if para.style is not None else ""
        paragraphs.append(
            {
                "index": index,
                "style": style_name,
                "text": text,
            }
        )

    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([_clean(cell.text) for cell in row.cells])
        tables.append({"index": table_index, "rows": rows})

    full_text = "\n\n".join(item["text"] for item in paragraphs)
    style_counter = Counter((p.get("style") or "Unknown") for p in paragraphs)

    return {
        "source_file": str(docx_path),
        "file_type": "docx",
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "full_text": full_text,
        "paragraphs": paragraphs,
        "tables": tables,
        "stats": {
            "full_text_chars": len(full_text),
            "style_counts": dict(style_counter),
        },
        "warnings": [],
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract text and tables from DOCX.")
    parser.add_argument("input", help="Path to input DOCX file")
    parser.add_argument("output", help="Path to output JSON file")
    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()

    if not input_path.exists():
        raise SystemExit(f"Input file not found: {input_path}")

    data = extract_docx_document(input_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote: {output_path}")


if __name__ == "__main__":
    main()
